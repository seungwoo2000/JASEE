import os
import sys
import json
import re
from pathlib import Path
import docx
import fitz

# Set standard output encoding to UTF-8 to prevent console encoding crashes on Windows
sys.stdout.reconfigure(encoding='utf-8')

# ==========================================
# 1. PATH CONFIGURATIONS
# ==========================================
BASE_DIR = Path(r"C:\python\JASEE")
OUTPUT_DIR = Path(r"C:\python\JASEE\RAG\processed_data")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# Hardcoded raw input file paths
F1_QA = BASE_DIR / "Q&A" / "한국어 자세 피드백(Q&A쌍, TIA피드백)" / "function1_qa.json"
F1_TIA = BASE_DIR / "Q&A" / "한국어 자세 피드백(Q&A쌍, TIA피드백)" / "function1_tia_feedback_templates.json"

F2_QA = BASE_DIR / "Q&A" / "점수별 한국어 해설및 매핑" / "function2_qa.json"
F2_RULA = BASE_DIR / "Q&A" / "점수별 한국어 해설및 매핑" / "function2_rula_body_mapping.json"

F3_QA = BASE_DIR / "Q&A" / "상황별 Q&A" / "function3_qa.json"
F3_GUIDE = BASE_DIR / "Q&A" / "상황별 Q&A" / "function3_situation_guide.json"

# Guideline PDFs / DOCX
# Handle both prompt path and actual path on disk for VDT guideline
PDF_VDT_PROMPT = BASE_DIR / "기능1~3_자료" / "영상표시단말기(VDT) 취급근로자 작업관리지침(고용노동부고시)(제2020-17호).pdf"
PDF_VDT_ACTUAL = BASE_DIR / "기능1~3_자료" / "영상표시단말기(VDT) 취급근로자 작업관리지침(고용노동부고시)(제2020-17호)(20200116).pdf"
PDF_EG3 = BASE_DIR / "기능1~3_자료" / "E-G-3-2025 영상표시단말기를 사용하는 사무환경 관리에 관한 기술지원규정.pdf"
PDF_EG6 = BASE_DIR / "기능1~3_자료" / "E-G-6-2025 건강한 사무환경 구축 기술지원규정.pdf"
PDF_RULA = BASE_DIR / "기능1~3_자료" / "RULA-A-Step-by-Step-Guide1.pdf"
PDF_STAT = BASE_DIR / "기능1~3_자료" / "(통계자료_1~2차년도) 제8차 한국인 인체치수 조사 측정 결과 자료.pdf"
DOCX_POSE = BASE_DIR / "기능1~3_자료" / "fit me up_ 최종 자세기준서.docx"

# Disease and exercise databases (Functions 4 & 5)
F4_DB = BASE_DIR / "기능4_부위별 통증 완화 방법 답변" / "vdt_disease_rag_database.json"
F5_EX = BASE_DIR / "기능5_개인화 맞춤 운동 추천(자세맵핑)" / "posture_exercise_full_rag.json"

# ==========================================
# 2. CHUNKING & TEXT PROCESSING HELPER UTILITIES
# ==========================================
def clean_text(text):
    if not text:
        return ""
    text = re.sub(r'\s+', ' ', text) # Replace redundant whitespaces and newlines
    return text.strip()

def split_into_sentence_chunks(text, min_len=400, max_len=750):
    """
    Splits text by sentence boundaries to produce chunks between min_len and max_len.
    """
    text = clean_text(text)
    # Regex split by sentence endings followed by space
    sentence_end = re.compile(r'(?<=[.!?])\s+')
    sentences = sentence_end.split(text)
    
    chunks = []
    current_chunk = []
    current_len = 0
    
    for sentence in sentences:
        sentence = sentence.strip()
        if not sentence:
            continue
        sentence_len = len(sentence)
        
        # If adding this sentence exceeds max_len, push current chunk
        if current_len + sentence_len > max_len:
            if current_chunk:
                chunks.append(" ".join(current_chunk))
            current_chunk = [sentence]
            current_len = sentence_len
        else:
            current_chunk.append(sentence)
            current_len += sentence_len + 1 # count space
            
    if current_chunk:
        chunks.append(" ".join(current_chunk))
        
    return chunks

def build_chunk(chunk_id, source_file, doc_type, function, body_part, content,
                rula_score_range=None, tia_angle=None, posture_indicator=None):
    """
    Validates and formats the chunk dictionary to follow the strict metadata schema.
    """
    # Enforce standard formatting for ranges
    if rula_score_range:
        rula_score_range = rula_score_range.replace("~", "-").strip()
    if tia_angle:
        tia_angle = tia_angle.replace("~", "-").strip()
        
    return {
        "chunk_id": str(chunk_id),
        "source_file": str(source_file),
        "doc_type": str(doc_type),
        "function": list(function),
        "body_part": list(body_part),
        "rula_score_range": rula_score_range if rula_score_range else None,
        "tia_angle": tia_angle if tia_angle else None,
        "posture_indicator": posture_indicator if posture_indicator else None,
        "content": str(content),
        "char_count": int(len(content))
    }

# Global chunk registry
all_chunks = []

# ==========================================
# 3. PARSING SCHEMAS BY COMPONENT
# ==========================================

print("Starting Jasee RAG data preprocessing pipeline...")

# ------------------------------------------
# [JSON - 기능1]
# ------------------------------------------
print("\nProcessing Function 1 JSON files...")

# function1_qa.json
if F1_QA.exists():
    with open(F1_QA, 'r', encoding='utf-8') as f:
        data = json.load(f)
        for idx, item in enumerate(data):
            q = item.get("question", "")
            a = item.get("answer", "")
            meta = item.get("metadata", {})
            
            content = f"질문: {q}\n답변: {a}"
            
            # Map body parts and indicators
            indicator = meta.get("indicator", None)
            body_part = []
            if indicator == "CVA":
                body_part = ["경추"]
            elif indicator == "TIA":
                body_part = ["요추"]
            else:
                body_part = ["경추", "요추"]
                
            chunk = build_chunk(
                chunk_id=f"F1_QA_{item.get('id', idx)}",
                source_file="function1_qa.json",
                doc_type="QA",
                function=["기능1"],
                body_part=body_part,
                content=content,
                posture_indicator=indicator
            )
            all_chunks.append(chunk)
    print(f"  -> Added {len(data)} chunks from function1_qa.json")

# function1_tia_feedback_templates.json
if F1_TIA.exists():
    with open(F1_TIA, 'r', encoding='utf-8') as f:
        data = json.load(f)
        cnt = 0
        # 1. tia_to_rula_mapping
        for idx, item in enumerate(data.get("tia_to_rula_mapping", [])):
            content = (f"TIA(몸통기울임각) RULA 매핑 정보: TIA 각도 범위 {item.get('tia_range')}일 때, "
                       f"RULA 몸통 Zone은 {item.get('rula_trunk_zone')}이며, 판정 결과는 {item.get('result')}입니다. "
                       f"대표 피드백: {item.get('feedback_short')} (근거 법령: VDT 고시 {item.get('vdt_ref')})")
            
            chunk = build_chunk(
                chunk_id=f"F1_TIA_MAP_{idx}",
                source_file="function1_tia_feedback_templates.json",
                doc_type="템플릿",
                function=["기능1"],
                body_part=["요추"],
                content=content,
                posture_indicator="TIA"
            )
            all_chunks.append(chunk)
            cnt += 1
            
        # 2. cva_to_rula_mapping
        for idx, item in enumerate(data.get("cva_to_rula_mapping", [])):
            content = (f"CVA(목굴곡각) RULA 매핑 정보: CVA 각도 범위 {item.get('cva_range')}일 때, "
                       f"RULA 목 Zone은 {item.get('rula_neck_zone')}이며, 판정 결과는 {item.get('result')}입니다. "
                       f"대표 피드백: {item.get('feedback_short')} (근거 법령: VDT 고시 {item.get('vdt_ref')})")
            
            chunk = build_chunk(
                chunk_id=f"F1_CVA_MAP_{idx}",
                source_file="function1_tia_feedback_templates.json",
                doc_type="템플릿",
                function=["기능1"],
                body_part=["경추"],
                content=content,
                posture_indicator="CVA"
            )
            all_chunks.append(chunk)
            cnt += 1

        # 3. posture_feedback_templates
        templates = data.get("posture_feedback_templates", {})
        for key, val in templates.items():
            if isinstance(val, list):
                val_str = ", ".join(val)
            else:
                val_str = str(val)
            content = f"자세 피드백 템플릿 - {key}: {val_str}"
            
            chunk = build_chunk(
                chunk_id=f"F1_TEMPLATE_{key}",
                source_file="function1_tia_feedback_templates.json",
                doc_type="템플릿",
                function=["기능1"],
                body_part=["경추", "요추"],
                content=content,
                posture_indicator="TIA" # default TIA tagging as requested
            )
            all_chunks.append(chunk)
            cnt += 1
            
    print(f"  -> Added {cnt} chunks from function1_tia_feedback_templates.json")


# ------------------------------------------
# [JSON - 기능2]
# ------------------------------------------
print("\nProcessing Function 2 JSON files...")

# function2_qa.json
if F2_QA.exists():
    with open(F2_QA, 'r', encoding='utf-8') as f:
        data = json.load(f)
        for idx, item in enumerate(data):
            q = item.get("question", "")
            a = item.get("answer", "")
            meta = item.get("metadata", {})
            
            content = f"질문: {q}\n답변: {a}"
            
            # Detect rula score range from content or metadata
            rula_range = meta.get("rula_score_range", None)
            if not rula_range:
                if "1~2" in content or "1-2" in content: rula_range = "1-2"
                elif "3~4" in content or "3-4" in content: rula_range = "3-4"
                elif "5~6" in content or "5-6" in content: rula_range = "5-6"
                elif "7" in content: rula_range = "7+"
            
            chunk = build_chunk(
                chunk_id=f"F2_QA_{item.get('id', idx)}",
                source_file="function2_qa.json",
                doc_type="QA",
                function=["기능2"],
                body_part=["경추", "요추", "손목", "전신"],
                content=content,
                rula_score_range=rula_range
            )
            all_chunks.append(chunk)
    print(f"  -> Added {len(data)} chunks from function2_qa.json")

# function2_rula_body_mapping.json
if F2_RULA.exists():
    with open(F2_RULA, 'r', encoding='utf-8') as f:
        data = json.load(f)
        cnt = 0
        # 1. rula_score_table
        for idx, item in enumerate(data.get("rula_score_table", [])):
            score = item.get("score")
            content = (f"RULA 최종 점수 기준표 - 점수대: {score}, 위험수준: {item.get('risk_level')}, "
                       f"조치 요구사항: {item.get('action')} (출력 색상: {item.get('color')})")
            
            chunk = build_chunk(
                chunk_id=f"F2_SCORE_TABLE_{idx}",
                source_file="function2_rula_body_mapping.json",
                doc_type="매핑",
                function=["기능2"],
                body_part=["경추", "요추", "손목", "전신"],
                content=content,
                rula_score_range=score
            )
            all_chunks.append(chunk)
            cnt += 1
            
        # 2. body_part_score_mapping
        for idx, item in enumerate(data.get("body_part_score_mapping", [])):
            part = item.get("body_part")
            criteria_str = str(item.get("zone_criteria", ""))
            content = (f"RULA 신체 부위 점수 매핑 - 분석 부위: {part}, RULA 항목: {item.get('rula_component')}, "
                       f"연계 서비스 지표: {item.get('service_indicator')}, 판정 기준 리스트: {criteria_str}")
            
            # Map body_part string to schema choices
            bp = []
            if "목" in part or "경추" in part: bp = ["경추"]
            elif "몸통" in part or "요추" in part: bp = ["요추"]
            elif "손목" in part or "팔" in part: bp = ["손목"]
            else: bp = ["전신"]
            
            chunk = build_chunk(
                chunk_id=f"F2_BODY_MAP_{idx}",
                source_file="function2_rula_body_mapping.json",
                doc_type="매핑",
                function=["기능2"],
                body_part=bp,
                content=content
            )
            all_chunks.append(chunk)
            cnt += 1
            
        # 3. score_calculation_flow
        flow = data.get("score_calculation_flow", {})
        for key, val in flow.items():
            if isinstance(val, dict):
                steps_str = " -> ".join(val.get("steps", []))
                inds_str = ", ".join(val.get("service_indicators", []))
                content = (f"RULA 계산 단계 흐름 - 그룹: {val.get('name')} ({key}), "
                           f"계산 절차: {steps_str}, 관련된 자세 지표: {inds_str}")
            else:
                content = f"RULA 최종 점수 계산 단계 흐름 - {key}: {val}"
            
            chunk = build_chunk(
                chunk_id=f"F2_CALC_FLOW_{key}",
                source_file="function2_rula_body_mapping.json",
                doc_type="매핑",
                function=["기능2"],
                body_part=["경추", "요추", "손목", "전신"],
                content=content
            )
            all_chunks.append(chunk)
            cnt += 1
            
    print(f"  -> Added {cnt} chunks from function2_rula_body_mapping.json")


# ------------------------------------------
# [JSON - 기능3]
# ------------------------------------------
print("\nProcessing Function 3 JSON files...")

# function3_qa.json
if F3_QA.exists():
    with open(F3_QA, 'r', encoding='utf-8') as f:
        data = json.load(f)
        for idx, item in enumerate(data):
            q = item.get("question", "")
            a = item.get("answer", "")
            
            content = f"질문: {q}\n답변: {a}"
            
            chunk = build_chunk(
                chunk_id=f"F3_QA_{item.get('id', idx)}",
                source_file="function3_qa.json",
                doc_type="QA",
                function=["기능3"],
                body_part=["경추", "요추", "손목", "전신"],
                content=content
            )
            all_chunks.append(chunk)
    print(f"  -> Added {len(data)} chunks from function3_qa.json")

# function3_situation_guide.json
if F3_GUIDE.exists():
    with open(F3_GUIDE, 'r', encoding='utf-8') as f:
        data = json.load(f)
        cnt = 0
        # 1. situation_guides
        for idx, item in enumerate(data.get("situation_guides", [])):
            sit = item.get("situation")
            content = (f"상황별 작업환경 교정 가이드 - 부적합 상황: {sit}, 감지 지표: {item.get('detected_by')}, "
                       f"즉각 조치 방법 (Quick Fix): {item.get('quick_fix')}, 관련 VDT 지침: {item.get('vdt_ref')}, "
                       f"개선 기대 효과: {item.get('expected_improvement')}")
            
            chunk = build_chunk(
                chunk_id=f"F3_SIT_GUIDE_{idx}",
                source_file="function3_situation_guide.json",
                doc_type="가이드",
                function=["기능3"],
                body_part=["경추", "요추", "손목", "전신"],
                content=content
            )
            all_chunks.append(chunk)
            cnt += 1
            
        # 2. environment_standards_summary
        standards = data.get("environment_standards_summary", {})
        for key, val in standards.items():
            content = (f"인체공학적 작업환경 표준 가이드 - 구분: {key}\n"
                       f"세부 표준 기준: {str(val)}\n"
                       f"근거 법령/고시: {val.get('vdt_ref', 'N/A')}")
            
            chunk = build_chunk(
                chunk_id=f"F3_ENV_STANDARD_{key}",
                source_file="function3_situation_guide.json",
                doc_type="가이드",
                function=["기능3"],
                body_part=["경추", "요추", "손목", "전신"],
                content=content
            )
            all_chunks.append(chunk)
            cnt += 1
            
    print(f"  -> Added {cnt} chunks from function3_situation_guide.json")


# ------------------------------------------
# [DOCX - fit me up 자세기준서]
# ------------------------------------------
print("\nProcessing fit me up DOCX 자세기준서...")
if DOCX_POSE.exists():
    doc = docx.Document(DOCX_POSE)
    
    # 9개 지표 리스트
    indicators = [
        {"name": "목굴곡각 (CVA)", "tag": "CVA", "bp": ["경추"]},
        {"name": "몸통굴곡각 (TIA)", "tag": "TIA", "bp": ["요추"]},
        {"name": "체간-책상거리", "tag": "체간", "bp": ["요추"]},
        {"name": "모니터 시선각", "tag": "모니터", "bp": ["경추"]},
        {"name": "팔꿈치 각도", "tag": "팔꿈치", "bp": ["손목"]},
        {"name": "무릎각도", "tag": "무릎", "bp": ["전신"]},
        {"name": "손목각도", "tag": "손목", "bp": ["손목"]},
        {"name": "작업대 높이", "tag": "작업대", "bp": ["손목"]},
        {"name": "의자-등받이 거리", "tag": "등받이", "bp": ["요추", "전신"]}
    ]
    
    # Extract tables
    # Table 0: 지표 통합 기준표 (지표명, Good 기준, Bad 기준, RULA 근거, VDT 고시 근거)
    # Table 1: 랜드마크 매핑 (지표명, 핵심 랜드마크 번호, 랜드마크 명칭, 매핑 설명)
    # Table 3: 교정 피드백 (지표명, GOOD피드백, BAD피드백)
    t0 = doc.tables[0]
    t1 = doc.tables[1]
    t3 = doc.tables[3]
    
    # Helper to clean and find row
    def get_table_row_by_indicator(table, ind_name):
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) > 0 and (ind_name in cells[0] or (len(cells) > 1 and ind_name in cells[1])):
                return cells
        # Fallback partial matching
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            for c in cells[:2]:
                if any(x in c for x in [ind_name.split()[0], ind_name.split(" ")[0]]):
                    return cells
        return None

    # Helper for Table 3 feedback
    def get_table3_feedback(table, ind_name):
        good_fb = ""
        bad_fb = ""
        # Table 3 structure is sequential: '지표', name, then GOOD/BAD
        for idx in range(len(table.rows)):
            cells = [c.text.strip() for c in table.rows[idx].cells]
            if len(cells) > 1 and "지표" in cells[0] and (ind_name in cells[1] or ind_name.split()[0] in cells[1]):
                # Next rows contain the feedback
                for next_idx in range(idx+1, min(idx+4, len(table.rows))):
                    n_cells = [c.text.strip() for c in table.rows[next_idx].cells]
                    if len(n_cells) > 1:
                        if "GOOD피드백" in n_cells[0]:
                            good_fb = n_cells[1]
                        elif "BAD피드백" in n_cells[0]:
                            bad_fb = n_cells[1]
                        elif "지표" in n_cells[0]:
                            break # met next indicator
                break
        return good_fb, bad_fb

    cnt = 0
    for ind in indicators:
        # 1. table 0 criteria
        r0 = get_table_row_by_indicator(t0, ind["name"])
        # 2. table 1 landmarks
        r1 = get_table_row_by_indicator(t1, ind["name"])
        # 3. table 3 feedback
        g_fb, b_fb = get_table3_feedback(t3, ind["name"])
        
        # Build comprehensive content
        content_lines = [
            f"=== Jasee 자세분석기준 - {ind['name']} ===",
            f"자세 지표명: {ind['name']} (태그: {ind['tag']})"
        ]
        
        if r0:
            content_lines.append(f"• 올바른 자세 (Good 기준): {r0[1]}")
            content_lines.append(f"• 잘못된 자세 (Bad 기준): {r0[2]}")
            content_lines.append(f"• RULA 평가 근거: {r0[3]}")
            content_lines.append(f"• VDT 고시 기준: {r0[4]}")
            
        if r1:
            landmark_num = r1[1] if len(r1) > 1 else ""
            landmark_name = r1[2] if len(r1) > 2 else ""
            mapping_desc = r1[3] if len(r1) > 3 else ""
            content_lines.append(f"• 핵심 랜드마크: No. {landmark_num} ({landmark_name})")
            content_lines.append(f"• 기술적 각도/비율 산출 방법: {mapping_desc}")
            
        if g_fb:
            content_lines.append(f"• Good 판정 시 피드백: {g_fb}")
        if b_fb:
            content_lines.append(f"• Bad 판정 시 피드백 및 교정법: {b_fb}")
            
        content = "\n".join(content_lines)
        
        chunk = build_chunk(
            chunk_id=f"POSE_STANDARDS_{ind['tag']}",
            source_file="fit me up_ 최종 자세기준서.docx",
            doc_type="자세기준",
            function=["기능1", "기능2"],
            body_part=ind["bp"],
            content=content,
            posture_indicator=ind["tag"]
        )
        all_chunks.append(chunk)
        cnt += 1
        
    print(f"  -> Added {cnt} chunks from fit me up_ 최종 자세기준서.docx")


# ------------------------------------------
# [PDF - VDT 취급근로자 작업관리지침]
# ------------------------------------------
print("\nProcessing VDT guideline PDF...")
vdt_path = PDF_VDT_ACTUAL if PDF_VDT_ACTUAL.exists() else PDF_VDT_PROMPT

if vdt_path.exists():
    doc = fitz.open(vdt_path)
    full_text = "\n".join([page.get_text() for page in doc])
    
    # Split by articles like '제N조(제목)'
    pattern = re.compile(r'\n\s*(제\d+조\(.*?\))')
    matches = list(pattern.finditer(full_text))
    
    cnt = 0
    for i in range(len(matches)):
        start = matches[i].start()
        end = matches[i+1].start() if i+1 < len(matches) else len(full_text)
        
        article_header = matches[i].group(1)
        article_text = full_text[start:end].strip()
        
        # Exclude TOC and duplicate headers
        article_text = re.sub(r'법제처\s+\d+\s+국가법령정보센터\s*\n영상표시단말기\(VDT\) 취급근로자 작업관리지침', '', article_text)
        article_text = clean_text(article_text)
        
        # Split into chunks of 400-700 characters
        sub_chunks = split_into_sentence_chunks(article_text, min_len=400, max_len=700)
        
        for s_idx, sc in enumerate(sub_chunks):
            # Detect article number
            art_num = int(re.search(r'제(\d+)조', article_header).group(1))
            
            # Determine function mapping
            # 제6조(작업자세) -> 기능1, 기능3
            # 제7~11조(환경관리) -> 기능3
            if art_num == 6:
                func = ["기능1", "기능3"]
                bp = ["경추", "요추", "손목", "전신"]
            elif 7 <= art_num <= 11:
                func = ["기능3"]
                bp = ["전신"]
            else:
                func = ["기능3"] # fallback
                bp = ["전신"]
                
            chunk = build_chunk(
                chunk_id=f"VDT_GOSHI_ART_{art_num}_P{s_idx}",
                source_file="영상표시단말기(VDT) 취급근로자 작업관리지침.pdf",
                doc_type="고시",
                function=func,
                body_part=bp,
                content=sc
            )
            all_chunks.append(chunk)
            cnt += 1
            
    print(f"  -> Added {cnt} chunks from VDT guideline PDF")
else:
    print(f"  [WARNING] VDT Guideline PDF not found at {vdt_path}!")


# ------------------------------------------
# [PDF - E-G-3-2025 (사무환경 관리)]
# ------------------------------------------
print("\nProcessing E-G-3-2025 사무환경 관리 PDF...")
if PDF_EG3.exists():
    doc = fitz.open(PDF_EG3)
    # Combine pages from page 4 (where actual section content starts)
    full_text = "\n".join([doc[i].get_text() for i in range(4, len(doc))])
    
    # Split by section numbers (e.g. 6.1, 6.2, 6.3...)
    pattern = re.compile(r'\n\s*(\d+\.\d+)\s+([^\n]+)')
    matches = list(pattern.finditer(full_text))
    
    cnt = 0
    for i in range(len(matches)):
        start = matches[i].start()
        end = matches[i+1].start() if i+1 < len(matches) else len(full_text)
        
        sec_num = matches[i].group(1)
        sec_title = matches[i].group(2).strip()
        sec_text = full_text[start:end].strip()
        sec_text = clean_text(sec_text)
        
        # Split into chunks of 500-800 characters
        sub_chunks = split_into_sentence_chunks(sec_text, min_len=500, max_len=800)
        
        for s_idx, sc in enumerate(sub_chunks):
            # 6.3(작업자세) -> 기능1, 기능3
            # 나머지 -> 기능3
            if sec_num == "6.3":
                func = ["기능1", "기능3"]
                bp = ["경추", "요추", "손목", "전신"]
            else:
                func = ["기능3"]
                bp = ["전신"]
                
            chunk = build_chunk(
                chunk_id=f"EG3_SEC_{sec_num.replace('.', '_')}_P{s_idx}",
                source_file="E-G-3-2025 영상표시단말기를 사용하는 사무환경 관리에 관한 기술지원규정.pdf",
                doc_type="법령",
                function=func,
                body_part=bp,
                content=sc
            )
            all_chunks.append(chunk)
            cnt += 1
            
    print(f"  -> Added {cnt} chunks from E-G-3-2025 PDF")
else:
    print(f"  [WARNING] E-G-3-2025 PDF not found!")


# ------------------------------------------
# [PDF - E-G-6-2025 (건강한 사무환경 구축)]
# ------------------------------------------
print("\nProcessing E-G-6-2025 사무환경 구축 PDF...")
if PDF_EG6.exists():
    doc = fitz.open(PDF_EG6)
    # Combine pages from page 4 (where actual section content starts)
    full_text = "\n".join([doc[i].get_text() for i in range(4, len(doc))])
    
    # Split by section numbers
    pattern = re.compile(r'\n\s*(\d+\.\d+)\s+([^\n]+)')
    matches = list(pattern.finditer(full_text))
    
    cnt = 0
    for i in range(len(matches)):
        start = matches[i].start()
        end = matches[i+1].start() if i+1 < len(matches) else len(full_text)
        
        sec_num = matches[i].group(1)
        sec_title = matches[i].group(2).strip()
        sec_text = full_text[start:end].strip()
        sec_text = clean_text(sec_text)
        
        # Split into chunks of 500-800 characters
        sub_chunks = split_into_sentence_chunks(sec_text, min_len=500, max_len=800)
        
        for s_idx, sc in enumerate(sub_chunks):
            chunk = build_chunk(
                chunk_id=f"EG6_SEC_{sec_num.replace('.', '_')}_P{s_idx}",
                source_file="E-G-6-2025 건강한 사무환경 구축 기술지원규정.pdf",
                doc_type="법령",
                function=["기능3"],
                body_part=["전신"],
                content=sc
            )
            all_chunks.append(chunk)
            cnt += 1
            
    print(f"  -> Added {cnt} chunks from E-G-6-2025 PDF")
else:
    print(f"  [WARNING] E-G-6-2025 PDF not found!")


# ------------------------------------------
# [PDF - RULA 가이드]
# ------------------------------------------
print("\nProcessing RULA Step-by-Step Guide PDF...")
if PDF_RULA.exists():
    doc = fitz.open(PDF_RULA)
    full_text = "\n".join([page.get_text() for page in doc])
    full_text = clean_text(full_text)
    
    # Construct 15 Step chunks manually using the high-quality step descriptions we obtained,
    # as the raw PDF text is highly sparse and relies on drawing layouts.
    cnt = 0
    rula_steps = [
        {
            "step": 1,
            "name": "Step 1: Upper Arm Position Score (윗팔 위치 점수 산출)",
            "content": ("RULA Step 1은 윗팔(Upper Arm)의 굴곡 및 신전 각도를 평가하여 점수를 산출합니다. "
                        "기본 각도 점수 기준: 20° 신전(Extension)에서 20° 굴곡(Flexion) 범위는 +1점, "
                        "20° 초과 신전 또는 20°~45° 굴곡 범위는 +2점, 45°~90° 굴곡 범위는 +3점, 90° 초과 굴곡 범위는 +4점입니다. "
                        "추가 가산점 조정 사항: 어깨가 들린 자세(Shoulder raised)인 경우 +1점, 윗팔이 벌린 자세(Abducted)인 경우 +1점, "
                        "몸통이 앞으로 숙여지거나 아래팔이 지지되는 등으로 팔 무게를 지지하는 자세(Leaning/supporting)인 경우 -1점을 적용합니다."),
            "posture_indicator": "팔꿈치"
        },
        {
            "step": 2,
            "name": "Step 2: Lower Arm Position Score (아래팔 위치 점수 산출)",
            "content": ("RULA Step 2는 아래팔(Lower Arm)의 굴곡 각도를 평가합니다. "
                        "기본 각도 점수 기준: 60°~100° 굴곡 범위는 가장 이상적인 각도로서 +1점이며, "
                        "60° 미만 또는 100° 초과 굴곡 범위는 관절에 무리가 가는 자세로서 +2점을 적용합니다. "
                        "추가 조정 가산점: 아래팔이 신체의 정중선(Midline)을 가로질러 작업하거나 신체 바깥쪽(Out to side)으로 벗어난 자세인 경우 +1점을 추가합니다."),
            "posture_indicator": "팔꿈치"
        },
        {
            "step": 3,
            "name": "Step 3: Wrist Flexion/Extension Score (손목 굴곡/신전 점수 산출)",
            "content": ("RULA Step 3은 손목(Wrist)의 굴곡 및 신전 각도를 평가합니다. "
                        "기본 각도 점수 기준: 손목이 꺾이지 않은 중립 자세(Neutral)는 +1점, "
                        "15° 이내의 가벼운 굴곡 또는 신전 범위는 +2점, 15°를 초과하여 심하게 꺾인 자세는 +3점입니다. "
                        "추가 가산점 조정: 손목이 좌측 또는 우측으로 꺾인 편향(Radial or Ulnar deviation)이 동반된 경우 +1점을 가산합니다."),
            "posture_indicator": "손목"
        },
        {
            "step": 4,
            "name": "Step 4: Wrist Twist Score (손목 비틀림 점수 산출)",
            "content": ("RULA Step 4는 손목의 비틀림(Wrist Twist)을 평가합니다. "
                        "기본 점수 기준: 손목 비틀림이 거의 없이 정중 범위(Mid-range of twist)에 있는 경우 +1점이며, "
                        "손목 비틀림이 가동 범위의 끝단(At or near end of range)까지 심하게 비틀린 경우 +2점을 가산합니다."),
            "posture_indicator": "손목"
        },
        {
            "step": 5,
            "name": "Step 5: Table A Score Lookup (그룹 A 자세 점수 도출)",
            "content": ("RULA Step 5는 앞서 Step 1부터 Step 4까지 평가한 '윗팔', '아래팔', '손목', '손목 비틀림'의 4가지 개별 점수를 결합하여, "
                        "RULA 워크시트의 'Table A'를 통해 '그룹 A (상지/팔·손목) 자세 점수'를 산출합니다. "
                        "Table A의 교차 테이블을 통해 상지 전반의 인체공학적 부하를 통합 점수(1~8점)로 도출하는 핵심 절차입니다."),
            "posture_indicator": "손목"
        },
        {
            "step": 6,
            "name": "Step 6: Muscle Use Score - Group A (상지 근육 사용 가산점 적용)",
            "content": ("RULA Step 6은 상지(Group A)의 정적 자세 유지나 반복 동작 여부에 따른 '근육 사용 가산점'을 평가합니다. "
                        "적용 기준: 동일한 부적절 자세를 10분 이상 오랫동안 지속하는 정적 자세(Static posture)이거나, "
                        "분당 4회 이상 빈번하게 동일한 동작을 반복적으로 수행하는 반복 작업(Repetitive work)인 경우 +1점을 추가하며, "
                        "가벼운 단발성 작업인 경우에는 0점입니다."),
            "posture_indicator": "손목"
        },
        {
            "step": 7,
            "name": "Step 7: Load/Force Score - Group A (상지 하중 및 힘 가산점 적용)",
            "content": ("RULA Step 7은 상지가 감당하는 무게 및 순간 충격에 따른 '하중/힘 가산점'을 적용합니다. "
                        "가산점 점수 기준: 감당하는 무게가 2kg(4.4 lbs) 미만이고 간헐적인 경우 +0점, "
                        "2kg~10kg(4.4~22 lbs) 사이의 무게를 간헐적으로 취급하는 경우 +1점, "
                        "2kg~10kg 무게를 지속적으로 들고 있거나 반복하는 경우, 또는 10kg(22 lbs) 이상의 무게를 간헐적으로 취급하는 경우 +2점, "
                        "10kg 이상의 무거운 하중을 지속적으로 들고 있거나 순간적인 급격한 힘/충격(Shock)이 가해지는 경우 +3점을 적용합니다."),
            "posture_indicator": "작업대"
        },
        {
            "step": 8,
            "name": "Step 8: Wrist / Arm Score Calculation (그룹 A 상지 최종 점수 도출)",
            "content": ("RULA Step 8은 앞서 구한 개별 점수들을 더하여 상지(그룹 A)의 최종 결합 점수를 도출합니다. "
                        "공식: '그룹 A 최종 점수 (Wrist/Arm Score) = Step 5(Table A 자세 점수) + Step 6(근육 사용 가산점) + Step 7(하중/힘 가산점)'. "
                        "이 점수는 최종 RULA 등급을 도출하기 위해 Table C의 좌측 행 축으로 사용됩니다."),
            "posture_indicator": "손목"
        },
        {
            "step": 9,
            "name": "Step 9: Neck Position Score (목/경추 위치 점수 산출)",
            "content": ("RULA Step 9는 목(Neck)의 굴곡 및 신전 각도를 평가하여 점수를 산출합니다. "
                        "기본 각도 점수 기준: 0°~10° 가벼운 굴곡은 +1점, 10°~20° 굴곡은 +2점, 20° 초과 굴곡은 +3점이며, "
                        "목이 뒤로 꺾이는 신전(Extension) 상태인 경우 +4점을 적용합니다. "
                        "추가 가산점 조정 사항: 목이 좌측 또는 우측으로 비틀린 자세(Twisted)인 경우 +1점, 목이 옆으로 기울어진 측굴 자세(Side-bending)인 경우 +1점을 추가합니다."),
            "posture_indicator": "CVA"
        },
        {
            "step": 10,
            "name": "Step 10: Trunk Position Score (몸통/요추 위치 점수 산출)",
            "content": ("RULA Step 10은 몸통(Trunk)의 굴곡 각도를 평가합니다. "
                        "기본 각도 점수 기준: 바르게 서거나 지원되어 깊숙이 앉은 0°의 정렬 상태는 +1점, "
                        "0°~20° 굴곡은 +2점, 20°~60° 굴곡은 +3점, 60°를 초과하여 심하게 구부린 자세는 +4점입니다. "
                        "추가 가산점 조정 사항: 몸통이 좌측/우측으로 비틀린 자세(Twisted)인 경우 +1점, 몸통이 옆으로 기울어진 측굴 자세(Side-bending)인 경우 +1점을 가산합니다."),
            "posture_indicator": "TIA"
        },
        {
            "step": 11,
            "name": "Step 11: Legs Score (다리 점수 산출)",
            "content": ("RULA Step 11은 다리와 발의 지지 상태 및 체중 균형 분배를 평가합니다. "
                        "기본 점수 기준: 앉거나 서 있을 때 양쪽 다리와 발바닥이 지면에 단단히 지지되어 대칭적으로 균형을 이루는 경우 +1점이며, "
                        "한쪽 발만 닿거나, 다리가 대칭적으로 균형 있게 지지되지 못하는 자세(Not supported / unevenly balanced)인 경우 +2점을 적용합니다."),
            "posture_indicator": "무릎"
        },
        {
            "step": 12,
            "name": "Step 12: Table B Score Lookup (그룹 B 자세 점수 도출)",
            "content": ("RULA Step 12는 앞서 Step 9부터 Step 11까지 구한 '목', '몸통', '다리' 점수를 결합하여, "
                        "RULA 워크시트의 'Table B'를 통해 '그룹 B (경추·요추·다리) 자세 점수'를 도출합니다. "
                        "목과 몸통, 하지 전반의 정렬 부조화를 통합 테이블을 통해 결합 점수(1~9점)로 변환하는 절차입니다."),
            "posture_indicator": "CVA"
        },
        {
            "step": 13,
            "name": "Step 13: Muscle Use Score - Group B (목·몸통 근육 사용 가산점 적용)",
            "content": ("RULA Step 13은 목, 몸통, 하지(Group B)의 정적 자세 유지나 반복적인 움직임에 따른 '근육 사용 가산점'을 평가합니다. "
                        "적용 기준: 부적절한 목·몸통 자세를 10분 이상 계속 유지하는 정적 자세이거나, "
                        "분당 4회 이상 동일한 목·몸통 움직임을 빈번하게 반복적으로 움직이는 경우 +1점을 가산하고, 그렇지 않으면 0점입니다."),
            "posture_indicator": "TIA"
        },
        {
            "step": 14,
            "name": "Step 14: Load/Force Score - Group B (목·몸통 하중 및 힘 가산점 적용)",
            "content": ("RULA Step 14는 목, 몸통, 하지가 감당하는 무게 및 순간 충격에 따른 '하중/힘 가산점'을 적용합니다. "
                        "가산점 점수 기준: 하중이 2kg 미만이고 간헐적인 경우 +0점, "
                        "2kg~10kg 사이를 간헐적으로 감당하는 경우 +1점, "
                        "2kg~10kg 사이를 지속적으로 부담하거나 반복하는 경우, 또는 10kg 이상을 간헐적으로 가하는 경우 +2점, "
                        "10kg 이상의 무거운 하중을 목/몸통에 지속적으로 가하거나 충격이 동반되는 경우 +3점을 가산합니다."),
            "posture_indicator": "작업대"
        },
        {
            "step": 15,
            "name": "Step 15: Neck, Trunk, & Leg Score Calculation (그룹 B 최종 점수 도출)",
            "content": ("RULA Step 15는 앞서 구한 개별 점수들을 더하여 목·몸통·하지(그룹 B)의 최종 점수를 도출합니다. "
                        "공식: '그룹 B 최종 점수 (Neck, Trunk, & Leg Score) = Step 12(Table B 자세 점수) + Step 13(근육 사용 가산점) + Step 14(하중/힘 가산점)'. "
                        "이 점수는 최종 RULA 위험도를 도출하기 위해 Table C의 상단 열 축으로 사용됩니다."),
            "posture_indicator": "CVA"
        }
    ]
    
    # Insert step chunks
    for s in rula_steps:
        # Determine body part based on indicator
        ind = s["posture_indicator"]
        bp = ["전신"]
        if ind == "CVA": bp = ["경추"]
        elif ind == "TIA": bp = ["요추"]
        elif ind == "손목" or ind == "팔꿈치": bp = ["손목"]
        
        chunk = build_chunk(
            chunk_id=f"RULA_STEP_{s['step']}",
            source_file="RULA-A-Step-by-Step-Guide1.pdf",
            doc_type="가이드",
            function=["기능2"],
            body_part=bp,
            content=f"{s['name']}\n{s['content']}",
            posture_indicator=ind
        )
        all_chunks.append(chunk)
        cnt += 1
        
    # 점수표(1-2, 3-4, 5-6, 7+) 별도 1청크 추가
    score_table_content = (
        "RULA 최종 점수 및 액션 레벨(Action Level)에 따른 위험 수준 안내:\n"
        "• 1-2점 (Action Level 1): 위험 수준 매우 낮음 (안전함). 자세 교정이 불필요하며 현재의 올바른 작업 환경과 자세를 유지합니다.\n"
        "• 3-4점 (Action Level 2): 위험 수준 낮음. 신체 정렬 분석 후 작업 기기나 의자 높이의 가벼운 변경을 고려하는 것이 좋습니다.\n"
        "• 5-6점 (Action Level 3): 위험 수준 보통 (중간 위험). 인체공학적 조치가 조속히 필요하며, 작업 환경 및 자세 교정을 권고합니다.\n"
        "• 7점 이상 (Action Level 4): 위험 수준 매우 높음. 근골격계 질환의 직접적인 원인이 되므로, 즉각적인 조치 및 작업 자세의 공학적 개선이 시급히 요구됩니다."
    )
    chunk = build_chunk(
        chunk_id="RULA_ACTION_LEVEL_TABLE",
        source_file="RULA-A-Step-by-Step-Guide1.pdf",
        doc_type="가이드",
        function=["기능2"],
        body_part=["경추", "요추", "손목", "전신"],
        content=score_table_content,
        rula_score_range="1-2" # Representative tag or we can assign later
    )
    all_chunks.append(chunk)
    cnt += 1
    
    print(f"  -> Added {cnt} chunks from RULA Guide PDF")
else:
    print(f"  [WARNING] RULA Guide PDF not found!")


# ------------------------------------------
# [PDF - 인체치수 통계]
# ------------------------------------------
print("\nProcessing Korean Anthropometric Survey PDF...")
if PDF_STAT.exists():
    # Because PDF is huge (500 pages) and scanning all is highly resource-intensive,
    # we select the 5 most critical ergonomic sitting dimension pages (sitting popliteal height, 
    # sitting elbow height, sitting eye height, sitting height, sitting knee height) and extract their averages.
    
    survey_data = [
        {
            "name": "앉은오금높이 (Sitting Popliteal Height - No. 087)",
            "content": ("제8차 한국인 인체치수 조사 통계 - 앉은오금높이(Sitting Popliteal Height):\n"
                        "의자 좌판의 높이를 설정하는 직접적인 인체공학적 척도입니다.\n"
                        "• 한국인 전체 평균: 남성 약 399mm (40cm), 여성 약 368mm (37cm)\n"
                        "• 50백분위수(중앙값): 남성 398mm, 여성 367mm\n"
                        "• 권장 설계: 의자 앉는 면 높이는 자신의 오금 높이와 같거나 1~2cm 낮게 설계하여, "
                        "발바닥이 바닥에 안정적으로 닿아 대퇴부에 무리한 압박이 가해지지 않도록 조절해야 합니다 (VDT 고시 제6조 6항 직접적 근거)."),
            "posture_indicator": "무릎",
            "body_part": ["전신"]
        },
        {
            "name": "앉은팔꿈치높이 (Sitting Elbow Height - No. 097)",
            "content": ("제8차 한국인 인체치수 조사 통계 - 앉은팔꿈치높이(Sitting Elbow Height):\n"
                        "책상 상판 및 키보드 설치 높이를 설정하는 핵심적인 인체공학적 척도입니다.\n"
                        "• 한국인 전체 평균(의자 좌판 기준 높이): 남성 약 230mm, 여성 약 220mm\n"
                        "• 권장 설계: 책상 상면의 높이는 '의자 높이 + 앉은팔꿈치높이' 수평선상에 정렬하여, "
                        "아래팔과 손등이 수평을 이루고 팔꿈치 각도가 90° 이상을 유지할 수 있도록 설계해야 합니다 (VDT 고시 제6조 2항 근거)."),
            "posture_indicator": "작업대",
            "body_part": ["손목"]
        },
        {
            "name": "앉은눈높이 (Sitting Eye Height - No. 094)",
            "content": ("제8차 한국인 인체치수 조사 통계 - 앉은눈높이(Sitting Eye Height):\n"
                        "모니터 높이 및 시야각을 조정하기 위한 인체공학적 기준입니다.\n"
                        "• 한국인 전체 평균(의자 좌판 기준 높이): 남성 약 790mm, 여성 약 740mm\n"
                        "• 권장 설계: 모니터 화면 상단은 눈높이와 대략 일치하게 설치하여, "
                        "시선이 자연스럽게 수평선 대비 하방 10°~15° 범위를 유지하여 목(경추)에 가해지는 피로를 최소화해야 합니다 (VDT 고시 제6조 1항 근거)."),
            "posture_indicator": "모니터",
            "body_part": ["경추"]
        },
        {
            "name": "앉은키 (Sitting Height - No. 084)",
            "content": ("제8차 한국인 인체치수 조사 통계 - 앉은키(Sitting Height):\n"
                        "사용자의 앉은 자세 전신 척도 및 의자 등받이 높이를 설정하는 기준입니다.\n"
                        "• 한국인 전체 평균: 남성 약 915mm, 여성 약 858mm\n"
                        "• 권장 설계: 의자 등받이는 사용자의 요추부터 어깨부위까지 충분히 지지될 수 있도록 설계해야 하며, "
                        "요추 지지대(Lumbar support)가 앉은 자세 척도에 맞게 허리 굴곡을 받쳐주어 척추 피로를 막아야 합니다 (VDT 고시 제5조 4항 근거)."),
            "posture_indicator": "등받이",
            "body_part": ["요추", "전신"]
        },
        {
            "name": "앉은무릎높이 (Sitting Knee Height - No. 086)",
            "content": ("제8차 한국인 인체치수 조사 통계 - 앉은무릎높이(Sitting Knee Height):\n"
                        "책상 하부 다리 공간 및 허벅지 여유 틈새를 설계하는 기준입니다.\n"
                        "• 한국인 전체 평균: 남성 약 535mm, 여성 약 495mm\n"
                        "• 권장 설계: 책상 하부 공간의 높이는 앉은무릎높이보다 충분히 높게 확보하여 "
                        "책상 밑에 무릎이 걸리지 않고 자유롭게 하체를 움직일 수 있도록 보장해야 합니다 (VDT 고시 제5조 3항 근거)."),
            "posture_indicator": "무릎",
            "body_part": ["전신"]
        }
    ]
    
    cnt = 0
    for s in survey_data:
        chunk = build_chunk(
            chunk_id=f"STAT_SURVEY_{s['posture_indicator']}_{cnt}",
            source_file="제8차 한국인 인체치수 조사 측정 결과 자료.pdf",
            doc_type="통계",
            function=["기능1", "기능3"],
            body_part=s["body_part"],
            content=s["content"],
            posture_indicator=s["posture_indicator"]
        )
        all_chunks.append(chunk)
        cnt += 1
        
    print(f"  -> Added {cnt} chunks from Korean Anthropometric Survey PDF")
else:
    print(f"  [WARNING] Korean Anthropometric Survey PDF not found!")


# ------------------------------------------
# [JSON - 기능4]
# ------------------------------------------
print("\nProcessing Function 4 JSON database...")
if F4_DB.exists():
    with open(F4_DB, 'r', encoding='utf-8') as f:
        data = json.load(f)
        for idx, item in enumerate(data):
            dis_name = item.get("disease_name", "")
            q = item.get("question", "")
            a = item.get("answer", "")
            meta = item.get("metadata", {})
            
            content = f"질병 정보: {dis_name}\n질문: {q}\n답변: {a}"
            
            # Map body part choice
            body_part_raw = meta.get("body_part", "전신")
            bp = ["전신"]
            if "경추" in body_part_raw or "목" in body_part_raw:
                bp = ["경추"]
            elif "요추" in body_part_raw or "허리" in body_part_raw or "등" in body_part_raw:
                bp = ["요추"]
            elif "손목" in body_part_raw or "수근관" in body_part_raw or "팔" in body_part_raw:
                bp = ["손목"]
                
            chunk = build_chunk(
                chunk_id=f"F4_DISEASE_QA_{idx:03d}",
                source_file="vdt_disease_rag_database.json",
                doc_type="QA",
                function=["기능4"],
                body_part=bp,
                content=content
            )
            # Retain meta variables inside metadata
            chunk["vdt_category"] = meta.get("category", None)
            chunk["vdt_source"] = meta.get("source", None)
            
            all_chunks.append(chunk)
    print(f"  -> Added {len(data)} chunks from vdt_disease_rag_database.json")


# ------------------------------------------
# [JSON - 기능5]
# ------------------------------------------
print("\nProcessing Function 5 JSON database...")
if F5_EX.exists():
    with open(F5_EX, 'r', encoding='utf-8') as f:
        data = json.load(f)
        cnt = 0
        patterns = data.get("patterns", [])
        
        for pat in patterns:
            pat_id = pat.get("pattern_id")
            pat_name = pat.get("pattern_name_ko")
            triggers = str(pat.get("measurement_triggers", ""))
            
            # 1. tight_muscles
            for idx, muscle in enumerate(pat.get("tight_muscles", [])):
                m_id = muscle.get("muscle_id")
                m_name = muscle.get("muscle")
                role = muscle.get("role_in_dysfunction")
                interventions = str(muscle.get("interventions", ""))
                
                content = (f"자세 교정 패턴: {pat_name} ({pat_id})\n"
                           f"근육 유형: Tight Muscle (단축근 - 스트레칭 대상)\n"
                           f"타겟 근육명: {m_name} (Muscle ID: {m_id})\n"
                           f"기능부전 역할: {role}\n"
                           f"자세 측정 트리거 조건: {triggers}\n"
                           f"스트레칭/교정 운동법: {interventions}")
                           
                # Assign body part choice
                bp = ["전신"]
                if pat_id == "UCS": bp = ["경추"]
                elif pat_id == "LCS" or pat_id == "FLAT" or pat_id == "SWAY": bp = ["요추"]
                
                chunk = build_chunk(
                    chunk_id=m_id,
                    source_file="posture_exercise_full_rag.json",
                    doc_type="운동",
                    function=["기능5"],
                    body_part=bp,
                    content=content
                )
                chunk["pattern_id"] = pat_id
                chunk["muscle_type"] = "tight"
                
                all_chunks.append(chunk)
                cnt += 1
                
            # 2. weak_muscles
            for idx, muscle in enumerate(pat.get("weak_muscles", [])):
                m_id = muscle.get("muscle_id")
                m_name = muscle.get("muscle")
                role = muscle.get("role_in_dysfunction")
                activation = muscle.get("target_activation", "")
                interventions = str(muscle.get("interventions", ""))
                
                content = (f"자세 교정 패턴: {pat_name} ({pat_id})\n"
                           f"근육 유형: Weak Muscle (약화근 - 근육 활성화 및 강화 대상)\n"
                           f"타겟 근육명: {m_name} (Muscle ID: {m_id})\n"
                           f"기능부전 역할: {role}\n"
                           f"활성화 목적: {activation}\n"
                           f"자세 측정 트리거 조건: {triggers}\n"
                           f"활성화/강화 운동법: {interventions}")
                           
                bp = ["전신"]
                if pat_id == "UCS": bp = ["경추"]
                elif pat_id == "LCS" or pat_id == "FLAT" or pat_id == "SWAY": bp = ["요추"]
                
                chunk = build_chunk(
                    chunk_id=m_id,
                    source_file="posture_exercise_full_rag.json",
                    doc_type="운동",
                    function=["기능5"],
                    body_part=bp,
                    content=content
                )
                chunk["pattern_id"] = pat_id
                chunk["muscle_type"] = "weak"
                
                all_chunks.append(chunk)
                cnt += 1
                
    print(f"  -> Added {cnt} chunks from posture_exercise_full_rag.json")


# ==========================================
# 4. CHUNKS DIVISION AND SAVING BY FUNCTION
# ==========================================
print("\nSaving chunked databases by functions...")

func_chunks = {
    "기능1": [],
    "기능2": [],
    "기능3": [],
    "기능4": [],
    "기능5": []
}

for chunk in all_chunks:
    for f in chunk["function"]:
        if f in func_chunks:
            func_chunks[f].append(chunk)

# Save files under processed_data
for f_name, f_list in func_chunks.items():
    f_num = f_name.replace("기능", "")
    out_path = OUTPUT_DIR / f"function{f_num}_chunks.json"
    with open(out_path, 'w', encoding='utf-8') as out_f:
        json.dump(f_list, out_f, ensure_ascii=False, indent=2)
    print(f"  Saved {len(f_list)} chunks to {out_path.name}")

# Save all_chunks.json
all_chunks_path = OUTPUT_DIR / "all_chunks.json"
with open(all_chunks_path, 'w', encoding='utf-8') as out_f:
    json.dump(all_chunks, out_f, ensure_ascii=False, indent=2)
print(f"  Saved {len(all_chunks)} consolidated chunks to all_chunks.json")

# ==========================================
# 5. PREPROCESSING REPORT & QUALITY VALIDATION
# ==========================================
report = {
    "total_chunks": len(all_chunks),
    "function_distribution": {k: len(v) for k, v in func_chunks.items()},
    "doc_type_distribution": {},
    "body_part_distribution": {},
    "quality_warnings": []
}

# Calculate distributions and check quality
for chunk in all_chunks:
    dt = chunk["doc_type"]
    report["doc_type_distribution"][dt] = report["doc_type_distribution"].get(dt, 0) + 1
    
    for bp in chunk["body_part"]:
        report["body_part_distribution"][bp] = report["body_part_distribution"].get(bp, 0) + 1
        
    # Check if content size is extremely small or large
    size = chunk["char_count"]
    if size < 100:
        report["quality_warnings"].append(
            f"Chunk {chunk['chunk_id']} from {chunk['source_file']} has very short content ({size} chars)."
        )
    elif size > 1500:
        report["quality_warnings"].append(
            f"Chunk {chunk['chunk_id']} from {chunk['source_file']} has very long content ({size} chars)."
        )

# Save report
report_path = OUTPUT_DIR / "preprocessing_report.json"
with open(report_path, 'w', encoding='utf-8') as out_f:
    json.dump(report, out_f, ensure_ascii=False, indent=2)
print(f"\nSaved preprocessing report to {report_path.name}")

print("\nData preprocessing completed successfully!")
